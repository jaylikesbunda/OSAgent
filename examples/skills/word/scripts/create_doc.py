import sys
import json
import os
from docx import Document


def main():
    args_json = os.environ.get("OSA_SKILL_ARGS_JSON", "{}")
    try:
        args = json.loads(args_json)
    except Exception:
        args = {}

    title = args.get("title", "Untitled")
    paragraphs_raw = args.get("paragraphs", "")
    paragraphs = (
        [p.strip() for p in paragraphs_raw.split(",") if p.strip()]
        if paragraphs_raw
        else []
    )
    filename = args.get("filename", "document.docx")

    output_dir = os.environ.get("OUTPUT_DIR", ".")
    output_path = os.path.join(output_dir, filename)

    doc = Document()
    doc.add_heading(title, 0)
    for para in paragraphs:
        if para:
            doc.add_paragraph(para)
    doc.save(output_path)

    result = {"path": output_path, "title": title, "paragraphs": len(paragraphs)}
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
